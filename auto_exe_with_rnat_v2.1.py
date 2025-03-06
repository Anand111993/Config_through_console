#in this version a placeholder and value is added in GUI and its functionality
#pulling requried details from show command like CDP, Description,Interfaces and getting value for each placeholder.


import sys
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QFileDialog, QLabel, QMessageBox
from PyQt5.QtGui import QPixmap
from PyQt5.QtWidgets import QTableWidgetItem,QComboBox
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Cm
from PIL import Image
import pandas as pd
import os
import datetime
import re
from openpyxl import load_workbook
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QFileDialog, QLabel, QTableWidget, QTableWidgetItem, QScrollArea

class DocxPlaceholderApp(QWidget):
    def __init__(self, folder_path=""):
        super().__init__()
        self.folder_path = folder_path
        self.docxFilePath = None  # Initialize to None or a default path
        # self.cdp_hostnames = None
        # self.arp_table = None
        # self.interface_details = None
        self.extractedData = {
        'cdp_hostnames': [],
        'arp_table': [],
        'interface_details': []
        }
        self.placeholders = {}
        self.statusLabel = QLabel(self) 
        self.statusLabel = {'setText': print}
        #print(f'Debug: Received folder path: {self.folder_path}')  # Debug statement
        self.initUI()

    def initUI(self):
        self.setWindowTitle('DOCX Magic Generator')
        self.setGeometry(300, 300, 650, 400)
        
        layout = QVBoxLayout()

        self.folderPathLabel = QLabel(f'Folder Path: {self.folder_path}', self)
        layout.addWidget(self.folderPathLabel)
        self.setLayout(layout)
        
        self.uploadDocxButton = QPushButton('Upload DOCX File', self)
        self.uploadDocxButton.clicked.connect(self.uploadDocx)
        layout.addWidget(self.uploadDocxButton)


        self.placeholderTable = QTableWidget(0, 2, self)
        self.placeholderTable.setHorizontalHeaderLabels(['Placeholder', 'Value'])
        self.placeholderTable.horizontalHeader().setStretchLastSection(True)
        self.placeholderTable.resizeColumnsToContents()
        layout.addWidget(self.placeholderTable)


        self.uploadTextButton = QPushButton('Upload Text File', self)
        self.uploadTextButton.clicked.connect(self.uploadText)
        layout.addWidget(self.uploadTextButton)
        
        self.downloadExcelButton = QPushButton('Download Placeholders to Excel', self)
        self.downloadExcelButton.clicked.connect(self.downloadPlaceholdersToExcel)
        layout.addWidget(self.downloadExcelButton)
        
        self.uploadExcelButton = QPushButton('Upload Excel File', self)
        self.uploadExcelButton.clicked.connect(self.uploadExcel)
        layout.addWidget(self.uploadExcelButton)
        
        self.generateDocxButton = QPushButton('Generate DOCX File', self)
        self.generateDocxButton.clicked.connect(self.generateDocx)
        layout.addWidget(self.generateDocxButton)
        
        # Populate the table with placeholder keys and combo boxes
        #self.populateTable()

        self.statusLabel = QLabel('', self)
        layout.addWidget(self.statusLabel)
        
        # Initialize buttons and connect signals
        #self.setupButtons(layout)

        self.docxFilePath = ''
        self.excelFilePath = ''
        self.textFilePath = ''
        
        self.setLayout(layout)
        
        # After setting up UI, check for text files and process if available
       # self.checkAndProcessTextFiles()
        #   # Ensure folder_path is valid and not empty before listing files
        # if self.folder_path and os.path.isdir(self.folder_path):
        #     text_files = [f for f in os.listdir(self.folder_path) if f.endswith('.txt')]
        #     if text_files:
        #         text_file_path = os.path.join(self.folder_path, text_files[0])
        #         self.processTextFile(text_file_path)
        #     else:
        #         self.statusLabel.setText('No text files found in the directory.')
        # else:
        #     self.statusLabel.setText('No valid folder path set. Please select a valid directory.')


    def checkAndProcessTextFiles(self):
        # Check for text files in the given folder path and process the first found
        if self.folder_path and os.path.isdir(self.folder_path):
            text_files = [f for f in os.listdir(self.folder_path) if f.endswith('.txt')]
            if text_files:
                text_file_path = os.path.join(self.folder_path, text_files[0])
                self.processTextFile(text_file_path)
            else:
                self.statusLabel.setText('No text files found in the directory.')
        else:
            self.statusLabel.setText('No valid folder path set. Please select a valid directory.')




    def uploadDocx(self):
        options = QFileDialog.Options()
        self.docxFilePath, _ = QFileDialog.getOpenFileName(self, "Upload DOCX File", "", "DOCX Files (*.docx);;All Files (*)", options=options)
        if self.docxFilePath:
            self.loadPlaceholders()
            self.statusLabel.setText(f'DOCX File Uploaded: {self.docxFilePath}')        
            self.checkAndProcessTextFiles() 

   

    def loadPlaceholders(self):
        if not self.docxFilePath:
            self.statusLabel.setText('Please upload a DOCX file first.')
            return
        
        doc = Document(self.docxFilePath)
        placeholders = set()
        placeholder_pattern = r'\{([^{}]+)\}'

        # Extract placeholders from paragraphs
        for paragraph in doc.paragraphs:
            placeholders.update(re.findall(placeholder_pattern, paragraph.text))

        # Extract placeholders from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(re.findall(placeholder_pattern, paragraph.text))

        self.placeholders = {p: '' for p in sorted(placeholders)}  # Store placeholders as dictionary
        self.updateTable()  # Update the table in the GUI
        print(f"Loaded {len(self.placeholders)} placeholders.")  # Debug print to check number of placeholders loaded

    def updateTable(self):
        self.placeholderTable.setRowCount(0)  # Reset the table rows
        for key in self.placeholders:
            row_position = self.placeholderTable.rowCount()
            self.placeholderTable.insertRow(row_position)
            self.placeholderTable.setItem(row_position, 0, QTableWidgetItem(key))
            self.placeholderTable.setItem(row_position, 1, QTableWidgetItem(self.placeholders[key]))


    def processTextFile(self, file_path):
        print(f'Processing text file: {file_path}')
        self.cdp_hostnames = self.extract_cdp_neighbors(file_path)
        self.arp_table = self.extract_arp_table(file_path)
        self.interface_details = self.extract_interface_details(file_path)
        self.updateTableWithExtractedData()  # Refresh the ComboBoxes with new data
    

   
    def uploadText(self):
        options = QFileDialog.Options()
        self.textFilePath, _ = QFileDialog.getOpenFileName(self, "Upload Text File", "", "Text Files (*.txt);;All Files (*)", options=options)
        if self.textFilePath:
            self.folder_path = ''  # Clear the automatically selected path
            self.statusLabel.setText(f'Text File Uploaded: {self.textFilePath}')
            self.processTextFile(self.textFilePath)
            #self.updateTableWithTextData()  # New function to update the table

    
    # def updateTableWithExtractedData(self):
    #     interface_info = self.extractedData.get('interface_details', [])
    #     for row in range(self.placeholderTable.rowCount()):
    #         key_item = self.placeholderTable.item(row, 0)
    #         if key_item:
    #             key = key_item.text().strip()
    #             normalized_key = key.lower().replace('-', '_').replace(' ', '_')

    #               # Handling for interface-specific placeholders, assuming they are formatted with a port number
    #             if 'downstream_port' in normalized_key or 'downstream_device' in normalized_key:
    #                 port_number_match = re.search(r'\d+', normalized_key)
    #                 if port_number_match:
    #                     port_index = int(port_number_match.group()) - 1  # Adjust for zero-based indexing
    #                     if 0 <= port_index < len(interface_info):
    #                         port_name, description = interface_info[port_index]
    #                         # Decide which column to update based on key
    #                         if 'description' in normalized_key:
    #                             self.placeholderTable.setItem(row, 1, QTableWidgetItem(description))
    #                         else:
    #                             self.placeholderTable.setItem(row, 1, QTableWidgetItem(port_name))
    #                     else:
    #                         self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
    #                 else:
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem("Invalid placeholder format"))


    #             # Decide whether to add a ComboBox or a simple QLabel based on key type
    #             elif 'sh_cdp_neigh' in normalized_key or 'sh_ip_arp' in normalized_key or 'sh_int_status' in normalized_key:
    #                 # Directly display data or a placeholder message
    #                 items = self.getListOfValuesForKey(key)
    #                 if items == ["No data available"]:
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
    #                 else:
    #                     data_str = ", ".join(items)  # Convert list of items to a single string
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem(data_str))
    #             else:
    #                 # Insert a ComboBox for other keys
    #                 items = self.getListOfValuesForKey(key)
    #                 self.addEditableComboBoxToCell(row, 1, items)



#in this method getting uplink port but not getting it's description.
    # def updateTableWithExtractedData(self):
    #     # Extracted data from previous methods
    #     local_ports = self.extractedData.get('local_ports', [])  # List of unique local port names from extract_cdp_neighbors
    #     interface_info = self.extractedData.get('interface_details', [])  # List of (interface, description) tuples from extract_interface_details

    #     for row in range(self.placeholderTable.rowCount()):
    #         key_item = self.placeholderTable.item(row, 0)
    #         if key_item:
    #             key = key_item.text().strip()
    #             normalized_key = key.lower().replace('-', '_').replace(' ', '_')

    #             # Check if key is for uplink_port or upstream_device_description
    #             if 'uplink_port' in normalized_key or 'upstream_device_description' in normalized_key:
    #                 port_number_match = re.search(r'\d+', normalized_key)
    #                 if port_number_match:
    #                     port_index = int(port_number_match.group()) - 1  # Adjust for zero-based indexing
    #                     if 0 <= port_index < len(local_ports):
    #                         local_port = local_ports[port_index]

    #                         # Debugging: Print the port being checked
    #                         print(f"Checking Local Port: {local_port} for Row: {row}")

    #                         # Find the description corresponding to the local port
    #                         description = "No description found"
    #                         for port, desc in interface_info:
    #                             print(f"Comparing {local_port} with {port}")  # Debugging: Print comparison details
    #                             if port == local_port:
    #                                 description = desc
    #                                 break

    #                         # Debugging: Print the matched description
    #                         print(f"Matched Description for {local_port}: {description}")

    #                         if 'upstream_device_description' in normalized_key:
    #                             # Update the upstream device description
    #                             print(f"Updating description for key: {normalized_key} with value: {description}")
    #                             self.placeholderTable.setItem(row, 1, QTableWidgetItem(description))
    #                         elif 'uplink_port' in normalized_key:
    #                             # Update the uplink port
    #                             print(f"Updating uplink port for key: {normalized_key} with value: {local_port}")
    #                             self.placeholderTable.setItem(row, 1, QTableWidgetItem(local_port))
    #                     else:
    #                         self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
    #                 else:
    #                     self.placeholderTa

    #             # Check if key is for uplink_port or upstream_device_description
    #             if 'uplink_port' in normalized_key or 'upstream_device_description' in normalized_key:
    #                 port_number_match = re.search(r'\d+', normalized_key)
    #                 if port_number_match:
    #                     port_index = int(port_number_match.group()) - 1  # Adjust for zero-based indexing
    #                     if 0 <= port_index < len(local_ports):
    #                         local_port = local_ports[port_index]

    #                         # Find the description corresponding to the local port
    #                         description = next((desc for port, desc in interface_info if port == local_port), "No description found")

    #                         if 'description' in normalized_key:
    #                             # Update the upstream device description
    #                             self.placeholderTable.setItem(row, 1, QTableWidgetItem(description))
    #                         else:
    #                             # Update the uplink port
    #                             self.placeholderTable.setItem(row, 1, QTableWidgetItem(local_port))
    #                     else:
    #                         self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
    #                 else:
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem("Invalid placeholder format"))

    #             # Handling downstream ports and devices if needed
    #             elif 'downstream_port' in normalized_key or 'downstream_device' in normalized_key:
    #                 port_number_match = re.search(r'\d+', normalized_key)
    #                 if port_number_match:
    #                     port_index = int(port_number_match.group()) - 1  # Adjust for zero-based indexing
    #                     if 0 <= port_index < len(interface_info):
    #                         port_name, description = interface_info[port_index]
    #                         # Decide which column to update based on key
    #                         if 'description' in normalized_key:
    #                             self.placeholderTable.setItem(row, 1, QTableWidgetItem(description))
    #                         else:
    #                             self.placeholderTable.setItem(row, 1, QTableWidgetItem(port_name))
    #                     else:
    #                         self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
    #                 else:
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem("Invalid placeholder format"))

                   
    #             # Existing logic for other keys
    #             elif 'sh_cdp_neigh' in normalized_key or 'sh_ip_arp' in normalized_key or 'sh_int_status' in normalized_key:
    #                 items = self.getListOfValuesForKey(key)
    #                 if items == ["No data available"]:
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
    #                 else:
    #                     data_str = ", ".join(items)  # Convert list of items to a single string
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem(data_str))
    #             else:
    #                 # Insert a ComboBox for other keys
    #                 items = self.getListOfValuesForKey(key)
    #                 self.addEditableComboBoxToCell(row, 1, items)



    # def updateTableWithExtractedData(self):
    #     # Extracted data from previous methods
    #     local_ports = self.extractedData.get('local_ports', [])  # List of unique local port names from extract_cdp_neighbors
    #     interface_info = self.extractedData.get('interface_details', [])  # List of (interface, description) tuples from extract_interface_details

    #     # Debugging: Print local ports and interface descriptions
    #     print("Local Ports Extracted:", local_ports)
    #     print("Interface Info Extracted:", interface_info)

    #     for row in range(self.placeholderTable.rowCount()):
    #         key_item = self.placeholderTable.item(row, 0)
    #         if key_item:
    #             key = key_item.text().strip()
    #             normalized_key = key.lower().replace('-', '_').replace(' ', '_')

    #             # Debugging: Print the current key being processed
    #             print(f"Processing key: {key}, Normalized key: {normalized_key}")

    #             # Check if key is for uplink_port
    #             if 'uplink_port' in normalized_key:
    #                 port_number_match = re.search(r'\d+', normalized_key)
    #                 if port_number_match:
    #                     port_index = int(port_number_match.group()) - 1  # Adjust for zero-based indexing
    #                     if 0 <= port_index < len(local_ports):
    #                         local_port = local_ports[port_index]

    #                         # Debugging: Print the port being checked
    #                         print(f"Checking Local Port: {local_port} for Row: {row}")

    #                         # Find the description corresponding to the local port
    #                         description = "No description found"
    #                         for port, desc in interface_info:
    #                             print(f"Comparing {local_port} with {port}")  # Debugging: Print comparison details
    #                             if port == local_port:
    #                                 description = desc
    #                                 break

    #                         # Update the uplink port
    #                         print(f"Updating uplink port for key: {normalized_key} with value: {local_port}")
    #                         self.placeholderTable.setItem(row, 1, QTableWidgetItem(local_port))
    #                     else:
    #                         self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
    #                 else:
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem("Invalid placeholder format"))

    #             # Updated condition for upstream_device_description with incremental numbers
    #             elif re.match(r'upstream_device\d+_description_port\d+', normalized_key):
    #                 print(f"Found 'upstream_deviceX_description_portX' in key: {normalized_key}")  # Debugging

    #                 # Extract the port number from the normalized key
    #                 port_number_match = re.search(r'\d+', normalized_key)
    #                 if port_number_match:
    #                     port_index = int(port_number_match.group()) - 1  # Adjust for zero-based indexing
    #                     if 0 <= port_index < len(local_ports):
    #                         local_port = local_ports[port_index]

    #                         # Debugging: Print the port being checked
    #                         print(f"Checking Local Port: {local_port} for Row: {row}")

    #                         # Find the description corresponding to the local port
    #                         description = "No description found"
    #                         for port, desc in interface_info:
    #                             print(f"Comparing {local_port} with {port}")  # Debugging: Print comparison details
    #                             if port == local_port:
    #                                 description = desc
    #                                 break

    #                         # Debugging: Print the matched description
    #                         print(f"Matched Description for {local_port}: {description}")

    #                         # Update the upstream device description
    #                         print(f"Updating description for key: {normalized_key} with value: {description}")
    #                         self.placeholderTable.setItem(row, 1, QTableWidgetItem(description))
    #                     else:
    #                         self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
    #                 else:
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem("Invalid placeholder format"))

    #             # Handling downstream ports and devices if needed
    #             elif 'downstream_port' in normalized_key or 'downstream_device' in normalized_key:
    #                 port_number_match = re.search(r'\d+', normalized_key)
    #                 if port_number_match:
    #                     port_index = int(port_number_match.group()) - 1  # Adjust for zero-based indexing
    #                     if 0 <= port_index < len(interface_info):
    #                         port_name, description = interface_info[port_index]
    #                         # Decide which column to update based on key
    #                         if 'description' in normalized_key:
    #                             self.placeholderTable.setItem(row, 1, QTableWidgetItem(description))
    #                         else:
    #                             self.placeholderTable.setItem(row, 1, QTableWidgetItem(port_name))
    #                     else:
    #                         self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
    #                 else:
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem("Invalid placeholder format"))

    #             # Existing logic for other keys
    #             elif 'sh_cdp_neigh' in normalized_key or 'sh_ip_arp' in normalized_key or 'sh_int_status' in normalized_key:
    #                 items = self.getListOfValuesForKey(key)
    #                 if items == ["No data available"]:
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
    #                 else:
    #                     data_str = ", ".join(items)  # Convert list of items to a single string
    #                     self.placeholderTable.setItem(row, 1, QTableWidgetItem(data_str))
    #             else:
    #                 # Insert a ComboBox for other keys
    #                 items = self.getListOfValuesForKey(key)
    #                 self.addEditableComboBoxToCell(row, 1, items)

    def updateTableWithExtractedData(self):
        # Extracted data from previous methods
        local_ports = self.extractedData.get('local_ports', [])  # List of unique local port names from extract_cdp_neighbors
        interface_info = self.extractedData.get('interface_details', [])  # List of (interface, description) tuples from extract_interface_details

        # Debugging: Print local ports and interface descriptions
        print("Local Ports Extracted:", local_ports)
        print("Interface Info Extracted:", interface_info)

        # Track processed interfaces
        processed_interfaces = set()

        for row in range(self.placeholderTable.rowCount()):
            key_item = self.placeholderTable.item(row, 0)
            if key_item:
                key = key_item.text().strip()
                normalized_key = key.lower().replace('-', '_').replace(' ', '_')

                # Debugging: Print the current key being processed
                print(f"Processing key: {key}, Normalized key: {normalized_key}")

                # Check if key is for uplink_port
                if 'uplink_port' in normalized_key:
                    port_number_match = re.search(r'\d+', normalized_key)
                    if port_number_match:
                        port_index = int(port_number_match.group()) - 1  # Adjust for zero-based indexing
                        if 0 <= port_index < len(local_ports):
                            local_port = local_ports[port_index]

                            # Debugging: Print the port being checked
                            print(f"Checking Local Port for uplink: {local_port} for Row: {row}")

                            # Find the description corresponding to the local port
                            description = "No description found"
                            for port, desc in interface_info:
                                print(f"Comparing {local_port} with {port}")  # Debugging: Print comparison details
                                if port == local_port:
                                    description = desc
                                    break

                            # Update the uplink port
                            print(f"Updating uplink port for key: {normalized_key} with value: {local_port}")
                            self.placeholderTable.setItem(row, 1, QTableWidgetItem(local_port))
                            processed_interfaces.add(local_port)  # Mark as processed
                            print(f"Processed Interfaces after uplink: {processed_interfaces}")  # Debugging
                        else:
                            self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
                    else:
                        self.placeholderTable.setItem(row, 1, QTableWidgetItem("Invalid placeholder format"))

                # Updated condition for upstream_device_description with incremental numbers
                elif re.match(r'upstream_device\d+_description_port\d+', normalized_key):
                    print(f"Found 'upstream_deviceX_description_portX' in key: {normalized_key}")  # Debugging

                    # Extract the port number from the normalized key
                    port_number_match = re.search(r'\d+', normalized_key)
                    if port_number_match:
                        port_index = int(port_number_match.group()) - 1  # Adjust for zero-based indexing
                        if 0 <= port_index < len(local_ports):
                            local_port = local_ports[port_index]

                            # Debugging: Print the port being checked
                            print(f"Checking Local Port for upstream: {local_port} for Row: {row}")

                            # Find the description corresponding to the local port
                            description = "No description found"
                            for port, desc in interface_info:
                                print(f"Comparing {local_port} with {port}")  # Debugging: Print comparison details
                                if port == local_port:
                                    description = desc
                                    break

                            # Debugging: Print the matched description
                            print(f"Matched Description for {local_port}: {description}")

                            # Update the upstream device description
                            print(f"Updating description for key: {normalized_key} with value: {description}")
                            self.placeholderTable.setItem(row, 1, QTableWidgetItem(description))
                            processed_interfaces.add(local_port)  # Mark as processed
                            print(f"Processed Interfaces after upstream: {processed_interfaces}")  # Debugging
                        else:
                            self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
                    else:
                        self.placeholderTable.setItem(row, 1, QTableWidgetItem("Invalid placeholder format"))

                # Handling downstream ports and devices if needed
                elif 'downstream_port' in normalized_key or 'downstream_device' in normalized_key:
                    # Only process if the interface has not been processed yet
                    port_number_match = re.search(r'\d+', normalized_key)
                    if port_number_match:
                        port_index = int(port_number_match.group()) - 1  # Adjust for zero-based indexing
                        if 0 <= port_index < len(interface_info):
                            port_name, description = interface_info[port_index]

                            # Debugging: Print the current state of processed interfaces
                            print(f"Processed Interfaces before downstream: {processed_interfaces}")

                            # Check if the port from interface_info has already been processed
                            if port_name in processed_interfaces:
                                print(f"Skipping {port_name} as it is already processed for uplink/upstream.")  # Debugging
                                continue  # Skip processing this port as it's already used

                            # Update downstream device or port if not processed
                            if 'description' in normalized_key:
                                self.placeholderTable.setItem(row, 1, QTableWidgetItem(description))
                            else:
                                self.placeholderTable.setItem(row, 1, QTableWidgetItem(port_name))

                        else:
                            self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
                    else:
                        self.placeholderTable.setItem(row, 1, QTableWidgetItem("Invalid placeholder format"))

                # Existing logic for other keys
                elif 'sh_cdp_neigh' in normalized_key or 'sh_ip_arp' in normalized_key or 'sh_int_status' in normalized_key:
                    items = self.getListOfValuesForKey(key)
                    if items == ["No data available"]:
                        self.placeholderTable.setItem(row, 1, QTableWidgetItem("No data available"))
                    else:
                        data_str = ", ".join(items)  # Convert list of items to a single string
                        self.placeholderTable.setItem(row, 1, QTableWidgetItem(data_str))
                else:
                    # Insert a ComboBox for other keys
                    items = self.getListOfValuesForKey(key)
                    self.addEditableComboBoxToCell(row, 1, items)







     
    def getListOfValuesForKey(self, key):
        # Normalize the key to be lower case and replace potential delimiters
        normalized_key = key.lower().replace('-', '_').replace(' ', '_')

        # Debug output to see the normalized key and what is being checked against
        #print(f"Normalized Key: {normalized_key}")

        if 'sh_cdp_neigh' in normalized_key:
            data_available = self.extractedData['cdp_hostnames'] if self.extractedData['cdp_hostnames'] else ["No data available"]
            #print(f"Checking 'sh_cdp_neigh': Data Found - {data_available}")
            return data_available
        elif 'sh_ip_arp' in normalized_key:
            data_available = self.extractedData['arp_table'] if self.extractedData['arp_table'] else ["No data available"]
            #print(f"Checking 'sh_ip_arp': Data Found - {data_available}")
            return data_available
        elif 'sh_int_status' in normalized_key or 'interface' in normalized_key:
            # Extract only the interface names, assuming they are stored as the first item in a tuple
            data_available = [detail[0] for detail in self.extractedData['interface_details']] if self.extractedData['interface_details'] else ["No data available"]
            #print(f"Checking 'sh_int_status' or 'interface': Data Found - {data_available}")
            return data_available
        else:
            # Provide all extracted data for keys that do not match the specific conditions
            all_data = self.extractedData['cdp_hostnames'] + self.extractedData['arp_table'] + [detail[0] for detail in self.extractedData['interface_details']]
            all_data = list(set(all_data))  # Remove duplicates if any
            if all_data:
                #print(f"Returning all extracted data for key: {key}.")
                return all_data
            else:
                # If no data was extracted, still allow manual entry
                #print(f"No data available for key: {key}. Returning 'Manual entry'.")
                return ["Manual entry"]




    def addEditableComboBoxToCell(self, row, column, items):
        comboBox = QComboBox()
        comboBox.setEditable(True)
        comboBox.addItems(items)
        comboBox.setCurrentIndex(-1)  # No default selection
        self.placeholderTable.setCellWidget(row, column, comboBox)
        comboBox.currentIndexChanged.connect(lambda index, r=row, c=column, cb=comboBox: self.onComboBoxEdited(r, c, cb))

    def onComboBoxEdited(self, row, column, comboBox):
        selected_text = comboBox.currentText()
        print(f"Edited to '{selected_text}' in row {row}, column {column}")


  
    def downloadPlaceholdersToExcel(self):
        if not self.docxFilePath:
            self.statusLabel.setText('Please upload a DOCX file first.')
            return

        # Collect data from QTableWidget
        placeholder_data = []
        for row in range(self.placeholderTable.rowCount()):
            key = self.placeholderTable.item(row, 0).text()  # Placeholder key from the first column
            cell_widget = self.placeholderTable.cellWidget(row, 1)
            
            if isinstance(cell_widget, QComboBox):
                # If it's a ComboBox, get the current text
                value = cell_widget.currentText()
            else:
                # Otherwise, get the text from the QTableWidgetItem
                item = self.placeholderTable.item(row, 1)
                value = item.text() if item else ""  # Ensure we handle None values

            placeholder_data.append({'Key': key, 'Value': value})

        df = pd.DataFrame(placeholder_data)  # DataFrame from collected data

        # Save the DataFrame to Excel
        self.save_excel_with_formatting(df)

    def save_excel_with_formatting(self, df):
        # Define the file name with a timestamp
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        excel_file_name = f"placeholder_{timestamp}.xlsx"
        excel_path = os.path.join(os.getcwd(), excel_file_name)
        df.to_excel(excel_path, index=False)

        # Optionally, you can format the Excel file using openpyxl if needed
        # (e.g., adjust column widths)
        from openpyxl import load_workbook
        wb = load_workbook(excel_path)
        ws = wb.active
        for col in ws.columns:
            max_length = max(len(str(cell.value)) for cell in col if cell.value is not None)
            adjusted_width = max_length + 2
            ws.column_dimensions[col[0].column_letter].width = adjusted_width
        wb.save(excel_path)

        # Notify user
        self.statusLabel.setText(f'Placeholders and values saved to Excel: {excel_path}')


    def extract_arp_table(self, file_path):
        with open(file_path, 'r') as file:
            content = file.read()

        # Regular expression to find IP addresses
        ip_regex = r'\b(?:\d{1,3}\.){3}\d{1,3}\b'

        # Extracting IP addresses assuming they are present after specific table headers
        arp_table_start = content.find('IP ARP Table for context default')
        if arp_table_start != -1:
            # Locate the end of the section if needed, for example by another header or end of content
            arp_table_end = content.find('####', arp_table_start)
            arp_table_content = content[arp_table_start:arp_table_end] if arp_table_end != -1 else content[arp_table_start:]
            arp_table = re.findall(ip_regex, arp_table_content)
            self.extractedData['arp_table'] = arp_table
        else:
            arp_table = []

        return arp_table

    # def extract_cdp_neighbors(self, file_path):
    #     with open(file_path, 'r') as file:
    #         content = file.read()

    #     # Find the start of the CDP neighbors section
    #     cdp_section_start = content.find('#### sh cdp neigh ####')
    #     if cdp_section_start != -1:
    #         # Find the end of the section if possible
    #         cdp_section_end = content.find('####', cdp_section_start + 21)
    #         cdp_section_content = content[cdp_section_start:cdp_section_end] if cdp_section_end != -1 else content[cdp_section_start:]

    #         # Updated regex to extract full hostnames
    #         cdp_hostnames = re.findall(r'\b(\S+)\.t-mobile.net', cdp_section_content, re.M | re.I)
    #         # Remove duplicates by converting the list to a set and back to a list
    #         unique_cdp_hostnames = list(set(cdp_hostnames))
    #         print("Extracted CDP Hostnames:", unique_cdp_hostnames)  # Debug print
    #         self.extractedData['cdp_hostnames'] = unique_cdp_hostnames
    #         return unique_cdp_hostnames
    #     return []

    def extract_cdp_neighbors(self, file_path):
        with open(file_path, 'r') as file:
            content = file.read()

        # Find the start of the CDP neighbors section
        cdp_section_start = content.find('#### sh cdp neigh ####')
        if cdp_section_start != -1:
            # Find the end of the section if possible
            cdp_section_end = content.find('####', cdp_section_start + 21)
            cdp_section_content = content[cdp_section_start:cdp_section_end] if cdp_section_end != -1 else content[cdp_section_start:]

            # Corrected regex to extract Device-ID and Local Intrfce
            cdp_entries = re.findall(r'(\S+)\.t-mobile\.net.*\n\s+(\S+)', cdp_section_content, re.I)

            # Separate Device IDs and Local Interfaces correctly
            device_ids = [entry[0].split('.')[0] for entry in cdp_entries]  # Extracting only the part before '.t-mobile.net'
            local_interfaces = [entry[1] for entry in cdp_entries]

            # Remove duplicates by converting the list to a set and back to a list
            unique_device_ids = list(set(device_ids))
            unique_local_interfaces = list(set(local_interfaces))

            print("Extracted Device IDs:", unique_device_ids)  # Debug print
            print("Extracted Local Interfaces:", unique_local_interfaces)  # Debug print

            # Store extracted data
            self.extractedData['cdp_hostnames'] = unique_device_ids
            self.extractedData['local_ports'] = unique_local_interfaces

            return unique_device_ids, unique_local_interfaces
        return [], []



    def extract_interface_details(self, file_path):
        with open(file_path, 'r') as file:
            content = file.read()

        # Initialize the data structures
        connected_interfaces = []
        description_dict = {}

        # Extract interfaces and their statuses
        int_status_start = content.find('#### sh int status | in connected ####')
        if int_status_start != -1:
            int_status_content = content[int_status_start:]
            connected_interfaces = re.findall(r'^(\S+)\s+.*?connected', int_status_content, re.M)
            
            # Removing unwanted entries
            connected_interfaces = [iface for iface in connected_interfaces if iface not in ['####', 'mgmt0']]

            self.extractedData['interface_details'] = connected_interfaces
            print("Connected Interfaces:", connected_interfaces)

        # Extract descriptions
        int_desc_start = content.find('#### sh int description  ####')
        if int_desc_start != -1:
            # Find the end of the section based on the next section header
            int_desc_end = content.find('####', int_desc_start + 30)
            int_desc_content = content[int_desc_start:int_desc_end] if int_desc_end != -1 else content[int_desc_start:]
            
            # Use a regex to extract the interface names and descriptions
            all_descriptions = re.findall(r'(\S+)\s+eth\s+(\d+G|\d+)\s+(.*)', int_desc_content, re.M)
            
            # Create a dictionary from the extracted data
            description_dict = {desc[0]: desc[2] for desc in all_descriptions if len(desc) > 2}
            #print("Interface Descriptions:", description_dict)

        # Update the extracted interface details with descriptions if available
        detailed_interface_info = []
        for interface in connected_interfaces:
            description = description_dict.get(interface, "--")  # Use '--' if no description is found
            detailed_interface_info.append((interface, description))

        self.extractedData['interface_details'] = detailed_interface_info
        return detailed_interface_info





 
    # def extract_interface_details(self, file_path):
    #     with open(file_path, 'r') as file:
    #         content = file.read()

        
    #      # Initialize the data structure
    #     connected_interfaces = []
    #     description_dict = {}

    #     # Extract interfaces and their statuses
    #     int_status_start = content.find('#### sh int status | in connected ####')
    #     if int_status_start != -1:
    #         int_status_content = content[int_status_start:]
    #         connected_interfaces = re.findall(r'^(\S+)\s+.*?connected', int_status_content, re.M)
    #         # Use a refined regex to specifically target interface lines
    #         #connected_interfaces = re.findall(r'^(\S+)\s+to\s+[^\s]+ connected', int_status_content, re.M)
            
    #         # Removing unwanted entries
    #         connected_interfaces = [iface for iface in connected_interfaces if iface not in ['####', 'mgmt0']]

    #         self.extractedData['interface_details'] = connected_interfaces
    #         print(connected_interfaces)
    #      # Find the start of the interface description section
    #     int_desc_start = content.find('#### sh int description  ####')
    #     if int_desc_start != -1:
    #         # Find the end of the section based on the next section header
    #         int_desc_end = content.find('####', int_desc_start + 30)
    #         int_desc_content = content[int_desc_start:int_desc_end] if int_desc_end != -1 else content[int_desc_start:]
            
    #         # Use a regex to extract the interface names and descriptions
    #         all_descriptions = re.findall(r'(\S+)\s+eth\s+(\d+G|\d+)\s+(.*)', int_desc_content, re.M)
    #         #print('All Descriptions:', all_descriptions)
            
    #         # Create a dictionary from the extracted data
    #         description_dict = {desc[0]: desc[2] for desc in all_descriptions if len(desc) > 2}
    #         #print("Interface Descriptions:", description_dict)

    #         # # Create a dictionary from the interface descriptions
    #         # description_dict = {desc[0]: desc[2] for desc in all_descriptions}
    #     # Update the extracted interface details with descriptions if available
    #     detailed_interface_info = []
    #     for interface in connected_interfaces:
    #         description = description_dict.get(interface, "--")  # Use '--' if no description is found
    #         detailed_interface_info.append((interface, description))
        
    #     self.extractedData['interface_details'] = detailed_interface_info
    #     #print (detailed_interface_info)
    #     return detailed_interface_info
        

    def update_placeholders_with_data(self, df):
        cdp_hostnames_str = '\n'.join(self.cdp_hostnames) if self.cdp_hostnames else ""
        arp_table_str = '\n'.join(self.arp_table) if self.arp_table else ""
        interface_details_str = '\n'.join(self.interface_details) if self.interface_details else ""

        for index, row in df.iterrows():
            if row['Key'].strip().lower() == 'sh_cdp_neigh':
                df.at[index, 'Value'] = cdp_hostnames_str
            elif row['Key'].strip().lower() == 'sh_ip_arp':
                df.at[index, 'Value'] = arp_table_str
            elif row['Key'].strip().lower() == 'sh_int_status_|_in_connected':  # Ensure the placeholder matches
                df.at[index, 'Value'] = interface_details_str

        return df

    def uploadExcel(self):
        options = QFileDialog.Options()
        self.excelFilePath, _ = QFileDialog.getOpenFileName(self, "Upload Excel File", "", "Excel Files (*.xlsx);;All Files (*)", options=options)
        if self.excelFilePath:
            self.statusLabel.setText(f'Excel File Uploaded: {self.excelFilePath}')
            self.populateTableFromExcel()

    def populateTableFromExcel(self):
  
        try:
            # Load the Excel file into a DataFrame
            df = pd.read_excel(self.excelFilePath)
            
            # Ensure the 'Key' and 'Value' columns exist
            if 'Key' in df.columns and 'Value' in df.columns:
                # Convert all data in the 'Value' column to strings to prevent type issues
                df['Value'] = df['Value'].astype(str)
                value_dict = df.set_index('Key')['Value'].to_dict()

                # Iterate over the rows of the table widget
                for row in range(self.placeholderTable.rowCount()):
                    key_item = self.placeholderTable.item(row, 0)
                    if key_item:
                        key = key_item.text()
                        # Check if the key from the table exists in the dictionary
                        if key in value_dict:
                            cell_widget = self.placeholderTable.cellWidget(row, 1)
                            if isinstance(cell_widget, QComboBox):
                                # If the ComboBox already has a value, skip updating
                                if cell_widget.currentIndex() == -1:
                                    index = cell_widget.findText(value_dict[key], Qt.MatchExactly | Qt.MatchCaseSensitive)
                                    if index >= 0:
                                        cell_widget.setCurrentIndex(index)
                                    else:
                                        cell_widget.addItem(value_dict[key])
                                        cell_widget.setCurrentIndex(cell_widget.count() - 1)
                            else:
                                # Set the value only if the cell is empty
                                current_value = self.placeholderTable.item(row, 1)
                                if not current_value or not current_value.text():
                                    self.placeholderTable.setItem(row, 1, QTableWidgetItem(value_dict[key]))
            else:
                self.statusLabel.setText('Excel file must contain columns "Key" and "Value".')
        except Exception as e:
            self.statusLabel.setText(f'Failed to load Excel file: {str(e)}')

    def delete_rows_based_on_placeholders(self, doc, data_dict):
        table_found = False

        # Attempt to locate the specific table
        for table in doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            if 'Device name' in headers and 'Downstream Device' in headers and 'Downstream Port' in headers:
                target_table = table
                table_found = True
                print('Target table found with headers:', headers)
                break

        if not table_found:
            self.statusLabel.setText("Specific table not found in the document.")
            return False

        # Placeholder logic extended to check and delete rows
        to_delete = []
        for i, row in enumerate(target_table.rows[1:], start=1):  # Skip header
            if len(row.cells) >= 3:
                # Assuming placeholder keys are formed as described and data_dict is correctly populated
                device_placeholder = f"downstream_device{i:02}_description_port{i:02}"
                port_placeholder = f"downstream_port{i:02}"

                # Fetch placeholder values or assume "No data available" if absent
                device_desc = data_dict.get(device_placeholder, "No data available").strip()
                port_desc = data_dict.get(port_placeholder, "No data available").strip()

                # Debug print to verify what is being checked
                print(f"Row {i}: Device - {device_desc}, Port - {port_desc}")

                if device_desc == "No data available" and port_desc == "No data available":
                    to_delete.append(i)

        print(f"Rows marked for deletion: {to_delete}")

        # Delete rows from the bottom up to avoid reindexing issues
        for index in reversed(to_delete):
            del target_table._element[index]

        return True






    def generateDocx(self):
        if not self.docxFilePath:
            self.statusLabel.setText('Please upload a DOCX file first.')
            return

        doc = Document(self.docxFilePath)

        # Collect data from QTableWidget
        data_dict = {}
        for row in range(self.placeholderTable.rowCount()):
            key = self.placeholderTable.item(row, 0).text()
            cell_widget = self.placeholderTable.cellWidget(row, 1)
            if isinstance(cell_widget, QComboBox):
                value = cell_widget.currentText()
            else:
                value_item = self.placeholderTable.item(row, 1)
                value = value_item.text() if value_item else ""
            data_dict[key] = value

        # if  self.delete_rows_based_on_placeholders(doc, data_dict):
        #     self.statusLabel.setText(' rows deleted or target table  found.')
        # else:    
        #     print('No rows deleted or target table not found.')
        #     return
        # Delete rows based on placeholders before making any replacements
        self.delete_rows_based_on_placeholders(doc, data_dict)

        def replace_text_in_paragraph(paragraph, placeholder, value):
            full_text = ''.join([run.text for run in paragraph.runs])
            if '{' + placeholder + '}' in full_text:
                if placeholder == 'sh_ip_arp' and isinstance(value, str):
                    value = value.replace(', ', '\n')  # Format ARP values as newline-separated
                new_text = full_text.replace('{' + placeholder + '}', value)
                for run in paragraph.runs:
                    run.text = ''
                paragraph.runs[0].text = new_text

        def replace_image_placeholder(paragraph, placeholder, image_path):
            full_text = ''.join([run.text for run in paragraph.runs])
            if '{' + placeholder + '}' in full_text:
                for run in paragraph.runs:
                    run.text = ''
                if os.path.exists(image_path):
                    run = paragraph.add_run()
                    with Image.open(image_path) as img:
                        img_width, img_height = img.size
                        aspect_ratio = img_height / img_width
                        max_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                        max_height = doc.sections[0].page_height - doc.sections[0].top_margin - doc.sections[0].bottom_margin
                        width_in_twips = img_width * 1440 / img.info.get('dpi', (72, 72))[0]
                        height_in_twips = width_in_twips * aspect_ratio
                        if width_in_twips > max_width:
                            width_in_twips = max_width
                            height_in_twips = width_in_twips * aspect_ratio
                        if height_in_twips > max_height:
                            height_in_twips = max_height
                            width_in_twips = height_in_twips / aspect_ratio
                        run.add_picture(image_path, width=Inches(width_in_twips / 1440), height=Inches(height_in_twips / 1440))

        # Process each paragraph for text or image replacements
        for paragraph in doc.paragraphs:
            for key, value in data_dict.items():
                replace_text_in_paragraph(paragraph, key, value)
                if isinstance(value, str) and value.endswith(('.png', '.jpg', '.jpeg')):
                    replace_image_placeholder(paragraph, key, value)

        # Process each cell in all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in data_dict.items():
                            replace_text_in_paragraph(paragraph, key, value)
                            if isinstance(value, str) and value.endswith(('.png', '.jpg', '.jpeg')):
                                replace_image_placeholder(paragraph, key, value)

        # Save the modified document
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        final_doc_file_name = f"final_doc_{timestamp}.docx"
        output_path = os.path.join(os.getcwd(), final_doc_file_name)
        doc.save(output_path)
        self.statusLabel.setText(f'Document saved to {output_path}')
        self.show_completion_message(output_path)



    def show_completion_message(self, output_path):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Information)
        msg.setWindowTitle("Process Complete")
        msg.setText(f'Document successfully saved to {output_path}')
        msg.setStandardButtons(QMessageBox.Ok)
        msg.exec_()        







if __name__ == '__main__':
    folder_path = os.getcwd()  # Set the folder path as the current working directory or any desired path
    app = QApplication(sys.argv)
    ex = DocxPlaceholderApp(folder_path)
    ex.show()
    sys.exit(app.exec_())
